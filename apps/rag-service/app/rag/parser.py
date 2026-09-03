import io
import re
from typing import BinaryIO


class DocumentParsingError(Exception):
    """Raised when parsing a document fails."""
    pass


def parse_document(file_bytes: bytes, filename: str) -> str:
    """Parse document bytes into markdown text according to file extension."""
    ext = ""
    if "." in filename:
        ext = filename.rsplit(".", 1)[-1].lower()

    if ext in ("md", "markdown", "txt", "text", "csv", "json"):
        return _parse_text(file_bytes)
    elif ext == "pdf":
        return _parse_pdf(file_bytes)
    elif ext in ("docx", "doc"):
        return _parse_docx(file_bytes)
    else:
        # Default fallback to utf-8 text
        return _parse_text(file_bytes)


def _parse_text(file_bytes: bytes) -> str:
    try:
        return file_bytes.decode("utf-8")
    except UnicodeDecodeError:
        return file_bytes.decode("utf-8", errors="replace")


def _parse_pdf(file_bytes: bytes) -> str:
    try:
        from pypdf import PdfReader
    except ImportError:
        raise DocumentParsingError("pypdf is not installed. Please install pypdf to parse PDF files.")

    try:
        reader = PdfReader(io.BytesIO(file_bytes))
        pages_text: list[str] = []
        for idx, page in enumerate(reader.pages):
            text = page.extract_text() or ""
            text = text.strip()
            if text:
                pages_text.append(f"<!-- Page {idx + 1} -->\n{text}")
        if not pages_text:
            return ""
        return "\n\n".join(pages_text)
    except Exception as e:
        raise DocumentParsingError(f"Failed to parse PDF document: {e}") from e


def _parse_docx(file_bytes: bytes) -> str:
    try:
        from docx import Document
    except ImportError:
        raise DocumentParsingError("python-docx is not installed. Please install python-docx to parse Word files.")

    try:
        doc = Document(io.BytesIO(file_bytes))
        elements: list[str] = []

        for p in doc.paragraphs:
            text = p.text.strip()
            if not text:
                continue
            style_name = (p.style.name or "").lower() if p.style else ""
            if "heading 1" in style_name:
                elements.append(f"# {text}")
            elif "heading 2" in style_name:
                elements.append(f"## {text}")
            elif "heading 3" in style_name:
                elements.append(f"### {text}")
            elif "heading" in style_name:
                elements.append(f"#### {text}")
            else:
                elements.append(text)

        # Parse tables
        for table in doc.tables:
            table_rows: list[str] = []
            for row_idx, row in enumerate(table.rows):
                cells = [re.sub(r"\s+", " ", cell.text.strip()) for cell in row.cells]
                table_rows.append("| " + " | ".join(cells) + " |")
                if row_idx == 0:
                    table_rows.append("| " + " | ".join(["---"] * len(cells)) + " |")
            if table_rows:
                elements.append("\n".join(table_rows))

        return "\n\n".join(elements)
    except Exception as e:
        raise DocumentParsingError(f"Failed to parse DOCX document: {e}") from e
